import os
import re
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from bs4 import BeautifulSoup
import base64
from PIL import Image
from io import BytesIO
import markdown
import logging
from docx.oxml.shared import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)

def iter_block_items(parent):
    """
    按顺序迭代文档中的所有段落和表格
    这样可以保留文档的原始顺序
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("不支持的父元素类型")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_image_paragraph_indices(doc):
    """获取文档中每个图片所在的段落索引及其上下文"""
    img_locations = {}
    img_contexts = {}
    paragraph_text_map = {}
    
    # 先建立段落索引到文本内容的映射
    for i, para in enumerate(doc.paragraphs):
        paragraph_text_map[i] = para.text
    
    # 收集文档中的所有图片关系ID
    rel_ids = {}
    for rel_id, rel in doc.part.rels.items():
        if rel.reltype == RT.IMAGE:
            rel_ids[rel_id] = rel.target_ref
    
    # 扫描所有段落，查找图片及其位置
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            # 检查运行对象中的XML元素是否包含图片
            if run._element.xml:
                for drawing in run._element.findall(".//"+qn("w:drawing")):
                    blip = drawing.find(".//"+qn("a:blip"))
                    if blip is not None:
                        embed_id = blip.get(qn("r:embed"))
                        if embed_id and embed_id in rel_ids:
                            target_ref = rel_ids[embed_id]
                            if target_ref not in img_locations:
                                img_locations[target_ref] = []
                            img_locations[target_ref].append(i)
                            
                            # 扩展上下文窗口到前后5个段落
                            # 收集图片上下文（前后段落的内容）
                            context_paragraphs_before = []
                            context_paragraphs_after = []
                            
                            # 收集前面5个段落
                            for j in range(max(0, i-5), i):
                                context_paragraphs_before.append(paragraph_text_map.get(j, ""))
                            
                            # 收集后面5个段落
                            for j in range(i+1, min(len(doc.paragraphs), i+6)):
                                context_paragraphs_after.append(paragraph_text_map.get(j, ""))
                            
                            img_contexts[target_ref] = {
                                'paragraph_index': i,
                                'context_paragraphs_before': context_paragraphs_before,
                                'current_paragraph': para.text,
                                'context_paragraphs_after': context_paragraphs_after,
                                # 添加字符级上下文（截取图片附近的文本）
                                'text_before_image': para.text[:para.text.find(run.text) + len(run.text)] if run.text in para.text else "",
                                'text_after_image': para.text[para.text.find(run.text) + len(run.text):] if run.text in para.text else ""
                            }
    
    return img_locations, img_contexts

def extract_images_from_docx(doc_path, output_dir):
    """从docx文件中提取图片并保存到指定目录，同时保留图片的上下文信息"""
    doc = docx.Document(doc_path)
    image_paths = []
    
    # 获取图片位置及上下文
    img_locations, img_contexts = get_image_paragraph_indices(doc)
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 提取图片
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            image_data = rel.target_part.blob
            
            # 尝试确定图片格式
            try:
                img = Image.open(BytesIO(image_data))
                ext = img.format.lower()
            except:
                ext = "png"  # 默认使用png
            
            image_filename = f"image_{image_count}.{ext}"
            image_path = os.path.join(output_dir, image_filename)
            
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            # 获取图片上下文信息（如果有）
            context_info = img_contexts.get(rel.target_ref, {})
            
            image_paths.append((
                rel.target_ref, 
                image_path, 
                image_count, 
                context_info
            ))
    
    return image_paths

def get_heading_level(para):
    """获取标题级别，如果不是标题则返回0"""
    # 如果段落没有样式，检查字体和格式特征
    if not para.style or not hasattr(para.style, 'name'):
        # 检查段落是否只包含粗体文本且较短（典型标题特征）
        if para.text and len(para.text.strip()) < 100:  # 标题一般不会太长
            is_all_bold = all(run.bold for run in para.runs if hasattr(run, 'bold') and run.text.strip())
            
            if is_all_bold and para.runs:
                # 根据字体大小判断标题级别
                first_run = para.runs[0]
                if hasattr(first_run, 'font') and hasattr(first_run.font, 'size'):
                    if first_run.font.size:
                        try:
                            # 尝试获取字体大小并转换为数值
                            size_str = str(first_run.font.size)
                            size = float(size_str.replace('pt', '').strip())
                            
                            if size >= 20: return 1
                            elif size >= 18: return 2
                            elif size >= 16: return 3
                            elif size >= 14: return 4
                            elif size >= 12 and is_all_bold: return 5
                        except (ValueError, AttributeError):
                            pass
    
    # 1. 优先检查样式名称（最可靠的方法）
    if para.style and hasattr(para.style, 'name'):
        style_name = para.style.name.lower()
        
        # 检查标准标题样式名称模式
        if 'heading' in style_name or '标题' in style_name or 'title' in style_name:
            # 提取数字部分确定级别
            for i in range(1, 7):  # 支持1-6级标题
                pattern = f'heading {i}|heading{i}|标题 {i}|标题{i}|h{i}'
                if re.search(pattern, style_name):
                    return i
            
            # 特殊情况处理：Title通常是主标题
            if style_name == 'title' or style_name == '标题':
                return 1
    
    # 2. 检查段落文本特征
    text = para.text.strip()
    if text:
        # Markdown风格标题检测
        if text.startswith('#'):
            # 计算开头的#数量
            level = 0
            for char in text:
                if char == '#':
                    level += 1
                else:
                    break
            if 1 <= level <= 6 and (level == len(text) or text[level:level+1].isspace()):
                return level
        
        # 检查全大写、关键词等标题特征
        is_likely_heading = (
            text.isupper() or  # 全大写文本
            re.search(r'\b[A-Z]{2,}\b', text) or  # 包含连续大写单词
            any(keyword in text for keyword in ['RAG', 'LLM', 'API', 'PDF', 'HTML'])  # 技术关键词
        )
        
        if is_likely_heading and len(text) < 100:
            # 检查字体样式判断级别
            has_large_font = False
            has_bold = False
            
            for run in para.runs:
                if hasattr(run, 'bold') and run.bold:
                    has_bold = True
                
                if hasattr(run, 'font') and hasattr(run.font, 'size') and run.font.size:
                    try:
                        size_str = str(run.font.size)
                        if 'pt' in size_str:
                            size = float(size_str.replace('pt', '').strip())
                            if size > 14:
                                has_large_font = True
                    except (ValueError, AttributeError):
                        pass
            
            if has_large_font and has_bold:
                return 2  # 默认为二级标题
            elif has_large_font:
                return 3
            elif has_bold:
                return 4
    
    # 3. 文本缩进和格式检查
    para_indent = 0
    if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'left_indent'):
        if para.paragraph_format.left_indent:
            try:
                indent_str = str(para.paragraph_format.left_indent)
                para_indent = float(indent_str.replace('pt', '').strip())
            except (ValueError, AttributeError):
                pass
    
    # 没有缩进且文本较短的粗体段落更可能是标题
    if para_indent == 0 and len(text) < 100:
        all_bold = all(hasattr(run, 'bold') and run.bold for run in para.runs if run.text.strip())
        if all_bold and para.runs:
            return 4
    
    return 0

def extract_toc(doc):
    """提取文档的目录结构"""
    toc = []
    
    for para in doc.paragraphs:
        level = get_heading_level(para)
        if level > 0:
            toc.append((level, para.text))
    
    return toc

def generate_toc_md(toc):
    """生成Markdown格式的目录"""
    if not toc:
        return ""
    
    md_toc = ["## 目录\n"]
    
    # 过滤掉空标题
    filtered_toc = [(level, text) for level, text in toc if text.strip()]
    
    for level, text in filtered_toc:
        # 创建适当的缩进
        indent = "  " * (level - 1)
        # 创建合法的锚链接（移除特殊字符）
        anchor = re.sub(r'[^\w\s-]', '', text).strip().lower().replace(' ', '-')
        # 确保锚链接不为空
        if not anchor:
            continue
        md_toc.append(f"{indent}- [{text}](#{anchor})")
    
    return "\n".join(md_toc) + "\n\n"

def is_list_item(para):
    """检查段落是否是列表项"""
    # 检查段落是否有列表样式
    if para.style and hasattr(para.style, 'name'):
        style_name = para.style.name.lower()
        if any(list_style in style_name for list_style in ['list', '列表', 'bullet', 'number', '编号']):
            return True
    
    # 检查段落格式是否有缩进（列表通常有缩进）
    if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'left_indent'):
        if para.paragraph_format.left_indent and not para.paragraph_format.first_line_indent:
            # 左缩进但没有首行缩进，典型的列表格式
            return True
    
    # 检查段落文本开头是否有列表标记
    text = para.text.strip()
    
    # 如果文本为空，不是列表项
    if not text:
        return False
    
    # 检查无序列表标记（支持各种常见无序列表符号）
    if re.match(r'^[-–—•◦○※＊*+>·]\s', text):
        return True
    
    # 检查带圆点的列表项（扩展支持）
    if re.match(r'^[●■□▪▫◆◇▶▷►▻]\s', text):
        return True
    
    # 检查有序列表标记（支持多种格式：1. 1) 1、(1) 等）
    if re.match(r'^(\d+|[a-zA-Z]|[ivxIVX]+)[.、)]\s', text) or re.match(r'^\(\d+\)\s', text):
        return True
    
    # 检查中文数字列表（一、二、三、等）
    if re.match(r'^[一二三四五六七八九十]+[、.]\s*', text):
        return True
    
    # 检查特殊情况：如果段落很短且紧跟着的段落是列表项，它可能是列表标题
    if len(text) < 30 and text.endswith((':', '：')):
        return False  # 这通常是列表的标题，不是列表项本身
    
    # 检查一些特殊列表格式（例如：- [ ]）
    if re.match(r'^[-*+]\s+\[\s?\]|\[\s?x\s?\]', text, re.IGNORECASE):  # 任务列表
        return True
    
    # 寻找列表的视觉特征（如符号后跟空格和文本）
    for run in para.runs:
        if run.text.strip() and any(symbol in run.text[:2] for symbol in ['•', '○', '■', '◦']):
            return True
    
    return False

def get_list_level(para):
    """获取列表的缩进级别"""
    level = 0
    
    # 从段落缩进值精确计算级别
    if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'left_indent'):
        left_indent = para.paragraph_format.left_indent
        if left_indent:
            try:
                # 转换缩进值为数字
                indent_str = str(left_indent)
                indent_value = float(indent_str.replace('pt', '').replace('cm', '').strip())
                
                # 根据缩进值确定级别（通常每级缩进约36pt或0.5-1cm）
                if 'cm' in indent_str:
                    level = max(level, int(indent_value / 0.7))  # 约0.7cm一级
                else:
                    level = max(level, int(indent_value / 36))  # 约36pt一级
            except (ValueError, AttributeError):
                pass
    
    # 从样式名称推断级别
    if para.style and hasattr(para.style, 'name'):
        style_name = para.style.name.lower()
        
        # 检查样式名称中是否包含级别信息
        for i in range(9, 0, -1):  # 检查9级到1级
            if f"level {i}" in style_name or f"级别 {i}" in style_name or f"{i}" in style_name:
                level = max(level, i - 1)  # 转为0-based级别
                break
    
    # 通过文本缩进和格式特征判断级别
    text = para.text
    
    # 匹配不同类型的列表标记，并捕获前导空格
    prefix_match = re.match(r'^(\s*)([-–—•◦○※＊*+>·]|\d+[.、)]|\([a-zA-Z0-9]+\)|\[[xX\s]\])\s', text)
    
    if prefix_match:
        # 通过前缀空白判断缩进级别
        prefix_spaces = prefix_match.group(1)
        space_level = len(prefix_spaces) // 2
        level = max(level, space_level)
    
    # 根据列表符号类型判断级别
    if any(marker in text[:5] for marker in ['•', '◦', '○', '▪', '▫']):
        symbol_indent = 0
        # 不同符号代表不同嵌套级别
        if '◦' in text[:5] or '○' in text[:5]:
            symbol_indent = 1  # 二级列表项
        elif '▪' in text[:5] or '▫' in text[:5] or '■' in text[:5] or '□' in text[:5]:
            symbol_indent = 2  # 三级列表项
        level = max(level, symbol_indent)
    
    # 检查是否有任务列表（如- [ ]），这些通常是独立级别
    if re.match(r'^\s*[-*+]\s+\[\s?\]|\[\s?x\s?\]', text, re.IGNORECASE):
        level = max(level, text.find('[') // 4)  # 根据方括号位置确定级别
    
    # 根据列表的标记类型推断嵌套级别
    if re.match(r'^\s*\d+[.、)]', text):  # 数字列表
        level = max(level, text.find(re.search(r'\d+', text).group(0)) // 4)
    elif re.match(r'^\s*[a-zA-Z][.、)]', text):  # 字母列表
        level = max(level, text.find(re.search(r'[a-zA-Z]', text).group(0)) // 4)
    elif re.match(r'^\s*\([a-zA-Z0-9]+\)', text):  # 带括号的列表
        level = max(level, text.find('(') // 4)
    
    return level

def is_code_block(para):
    """检查段落是否是代码块"""
    # 1. 首先检查段落样式名称
    if para.style and hasattr(para.style, 'name'):
        style_name = para.style.name.lower()
        if any(code_style in style_name for code_style in [
            'code', '代码', 'verbatim', 'preformatted', 'source', 'program', 'command', 'terminal'
        ]):
            return True
    
    # 获取段落文本
    text = para.text.strip()
    if not text:  # 空文本不是代码块
        return False
    
    # 2. 检查字体特征（等宽字体是代码的典型特征）
    has_monospace_font = False
    monospace_fonts = ['courier', 'consolas', 'monaco', 'monospace', 'menlo', 'lucida console', 'dejavu sans mono', 'fixedsys']
    
    for run in para.runs:
        if hasattr(run, 'font') and hasattr(run.font, 'name'):
            if run.font.name and any(font in run.font.name.lower() for font in monospace_fonts):
                has_monospace_font = True
                break
    
    # 3. 检查背景颜色和阴影（代码块常有灰色或其他背景色）
    has_code_background = False
    
    # 检查段落背景
    if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'shading'):
        if para.paragraph_format.shading:
            has_code_background = True
    
    # 检查文本运行对象的背景色
    background_runs_ratio = 0.0
    total_text_length = len(text)
    text_with_background = 0
    
    for run in para.runs:
        run_text = run.text
        has_bg = False
        
        # 检查高亮色
        if hasattr(run, 'font') and hasattr(run.font, 'highlight_color') and run.font.highlight_color:
            has_bg = True
        
        # 检查背景填充
        if not has_bg and hasattr(run, '_element') and hasattr(run._element, 'rPr'):
            rPr = run._element.rPr
            if rPr is not None and hasattr(rPr, 'xpath'):
                try:
                    shading = rPr.xpath('./w:shd')
                    if shading and len(shading) > 0:
                        has_bg = True
                except:
                    pass
        
        # 检查直接在run元素上的背景属性
        if not has_bg and hasattr(run, '_element') and hasattr(run._element, 'get_or_add_rPr'):
            try:
                shade_elm = run._element.get_or_add_rPr().find(qn('w:shd'))
                if shade_elm is not None:
                    has_bg = True
            except:
                pass
        
        if has_bg:
            text_with_background += len(run_text)
    
    if total_text_length > 0:
        background_runs_ratio = text_with_background / total_text_length
    
    # 如果超过60%的文本有背景色，认为是代码块
    if background_runs_ratio > 0.6:
        has_code_background = True
    
    # 4. 检查缩进和格式特征
    has_code_indent = False
    if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'left_indent'):
        if para.paragraph_format.left_indent:
            try:
                indent_value = float(str(para.paragraph_format.left_indent).replace('pt', '').replace('cm', '').strip())
                if indent_value > 10:  # 通常代码块有较大缩进
                    has_code_indent = True
            except (ValueError, AttributeError):
                pass
    
    # 5. 检查代码特征 - 语法特征
    code_patterns = [
        # 常见编程语句开头
        r'^(function|def|class|import|from|var|let|const)\s+\w+',  # 函数、类、变量声明
        r'^(public|private|protected)\s+\w+\s+\w+',  # 访问修饰符
        r'^\s*(if|for|while|switch|try|catch)\s*\(',  # 控制结构
        r'^\s*return\s+.+;?\s*$',  # return语句
        r'^(SELECT|INSERT|UPDATE|DELETE|CREATE|ALTER)\s+',  # SQL
        
        # 标记和特殊语法
        r'<\?php|\?>',  # PHP标记
        r'^```\w*$',  # Markdown代码块标记
        r'^#!\/bin\/(bash|sh|python|perl)',  # Shebang行
        
        # 特殊符号和结构
        r'^\s*[{}]\s*$',  # 单独一行的花括号
        r'^\s*[\[\]]\s*$',  # 单独一行的方括号
        r'^(\s*)[\w\-]+:\s*\w+',  # YAML键值对
        
        # 命令行
        r'^(\$|>)\s+[\w\-\.]+',  # 命令行提示符
        r'\bgit\s+(commit|push|pull|clone|checkout|add)\b',  # Git命令
        r'\bdocker\s+(run|build|exec|ps|images)\b',  # Docker命令
        r'\bnpm\s+(install|run|build|start)\b',  # npm命令
        
        # 新增匹配模式
        r'curl\s+https?:\/\/',  # curl命令
        r'wget\s+https?:\/\/',  # wget命令
        r'ssh\s+\w+@[\w\.]+',   # ssh命令
        r'cd\s+[\w\/\-\.]+',    # cd命令
        r'pip\s+install',       # pip命令 
        r'apt\s+(install|update|upgrade)',  # apt命令
        r'yum\s+(install|update)',  # yum命令
    ]
    
    has_code_syntax = any(re.search(pattern, text) for pattern in code_patterns)
    
    # 其他代码相关特征
    has_code_markers = (
        '{' in text and '}' in text or  # 花括号（代码块）
        '=' in text and ';' in text or  # 赋值和语句结束
        ':' in text and any(kw in text for kw in ['if', 'else', 'for', 'while', 'try', 'except', 'finally']) or  # Python代码
        '<' in text and '>' in text and '/' in text  # HTML/XML标记
    )
    
    # 6. 检查是否是bash命令行代码（经常有灰色背景）
    is_bash_command = False
    bash_patterns = [
        r'^\s*cd\s+',
        r'^\s*mkdir\s+',
        r'^\s*ls\s+',
        r'^\s*rm\s+',
        r'^\s*sudo\s+',
        r'^\s*apt\s+',
        r'^\s*yum\s+',
        r'^\s*docker\s+',
        r'^\s*git\s+',
        r'^\s*npm\s+',
        r'^\s*python\s+',
        r'^\s*pip\s+',
        r'^\s*javac\s+',
        r'^\s*mv\s+',
        r'^\s*cp\s+',
        r'^http',
        r'^https',
    ]
    
    for pattern in bash_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            is_bash_command = True
            break
    
    # 7. 综合判断
    # 如果有明显等宽字体或代码背景，基本确定是代码
    if (has_monospace_font and (has_code_background or has_code_syntax)) or \
       (has_code_background and has_code_syntax) or \
       has_code_background:  # 只要有背景色就认为是代码块
        return True
    
    # 如果有缩进和代码标记，很可能是代码
    if has_code_indent and (has_code_syntax or has_code_markers):
        return True
    
    # 如果是命令行代码
    if is_bash_command and (has_code_background or has_monospace_font):
        return True
    
    # 如果有明显的代码特征且不是很长的文本（避免误判长文本）
    if has_code_syntax and len(text) < 200:
        return True
    
    # 检测是否包含URL、文件路径等技术内容但不是简单的网址
    if any(tech_term in text for tech_term in ['http://', 'https://', 'file://', '/usr/', './']) and \
       not text.startswith(('http://', 'https://')) and len(text) < 150:
        return True
    
    # 行内命令检测，通常比较短且有特定格式
    if len(text) < 80 and any(cmd in text.lower() for cmd in [
        'cd ', 'mkdir', 'rmdir', 'touch', 'chmod', 'mv ', 'cp ', 'rm ', 'tar ', 'zip ', 'unzip', 'ping ', 'curl '
    ]):
        return True
    
    return False

def format_list_item(para):
    """将段落格式化为Markdown列表项"""
    text = para.text.strip()
    level = get_list_level(para)
    indent = '  ' * level
    
    # 处理有序列表 (数字格式)
    num_match = re.match(r'^(\d+)([.、)）])\s*(.*)', text)
    if num_match:
        list_num = num_match.group(1)
        list_text = num_match.group(3).strip()
        return f"{indent}{list_num}. {list_text}"
    
    # 处理有序列表 (字母格式)
    alpha_match = re.match(r'^([a-zA-Z])([.、)）])\s*(.*)', text)
    if alpha_match:
        list_char = alpha_match.group(1)
        list_text = alpha_match.group(3).strip()
        # 保留原始字母标记风格，但转换为Markdown格式
        return f"{indent}1. {list_text} (原标记: {list_char})"
    
    # 处理有序列表 (带括号的格式，如(1))
    paren_match = re.match(r'^\((\d+|[a-zA-Z])\)\s*(.*)', text)
    if paren_match:
        list_num = paren_match.group(1)
        list_text = paren_match.group(2).strip()
        # 尝试将字母转为数字，数字保持原样
        if list_num.isdigit():
            return f"{indent}{list_num}. {list_text}"
        else:
            return f"{indent}1. {list_text} (原标记: ({list_num}))"
    
    # 处理中文数字列表（一、二、三、等）
    cn_num_match = re.match(r'^([一二三四五六七八九十]+)([、.]\s*)(.*)', text)
    if cn_num_match:
        # 中文数字映射
        cn_nums = {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, '七':7, '八':8, '九':9, '十':10}
        cn_num = cn_num_match.group(1)
        list_text = cn_num_match.group(3).strip()
        
        # 尝试转换中文数字为阿拉伯数字
        if cn_num in cn_nums:
            return f"{indent}{cn_nums[cn_num]}. {list_text}"
        else:
            return f"{indent}1. {list_text} (原标记: {cn_num}、)"
    
    # 处理带圆点的无序列表
    circle_match = re.match(r'^([●■□▪▫◆◇▶▷►▻])\s*(.*)', text)
    if circle_match:
        list_text = circle_match.group(2).strip()
        return f"{indent}- {list_text}"
    
    # 处理任务列表
    task_match = re.match(r'^[-*+]\s+\[\s?\]\s*(.*)', text)
    if task_match:
        list_text = task_match.group(1).strip()
        return f"{indent}- [ ] {list_text}"
    
    task_checked_match = re.match(r'^[-*+]\s+\[[xX]\]\s*(.*)', text)
    if task_checked_match:
        list_text = task_checked_match.group(1).strip()
        return f"{indent}- [x] {list_text}"
    
    # 处理无序列表（检测各种无序列表标记）
    unordered_match = re.match(r'^([-–—•◦○※＊*+>·])\s*(.*)', text)
    if unordered_match:
        list_text = unordered_match.group(2).strip()
        return f"{indent}- {list_text}"
    
    # 如果以上模式都不匹配但段落有列表样式，则作为无序列表处理
    if para.style and hasattr(para.style, 'name'):
        style_name = para.style.name.lower()
        if any(list_style in style_name for list_style in ['list', '列表', 'bullet']):
            # 检查样式名来确定是有序还是无序列表
            if any(num_style in style_name for num_style in ['number', '编号', 'order']):
                return f"{indent}1. {text}"
            else:
                return f"{indent}- {text}" 
    
    # 处理其他不明确的情况，默认作为无序列表项
    return f"{indent}- {text}"

def format_paragraph(para):
    """格式化段落为Markdown"""
    # 如果段落为空，返回空行
    if not para.text.strip():
        return ""
    
    # 将段落文本复制一份以便修改
    text = para.text
    
    # 保留特殊符号 - 扩展列表
    special_symbols = [
        # 表情符号
        '❓', '❗', '✅', '✓', '✔️', '✗', '✘', '★', '☆', '➔', '➤', '➡️', '⬅️', '⬆️', '⬇️',
        '📌', '📝', '📊', '📈', '📉', '📋', '⚠️', '⚡', '🔍', '🔎', '🔑', '🔒', '🔓', '💡',
        '📖', '📚', '📄', '📃', '📜', '📰', '🗂️', '📁', '📂', '📤', '📥', '📏', '📐', '📞',
        '⌨️', '🖥️', '🔆', '🔅', '⏱️', '⏰', '🖋️', '✒️', '🖊️', '📱', '📶', '🚩', '⏯️',
        
        # 常用数学和科学符号
        '∑', '∏', '∫', '∂', '∇', '√', '∛', '∞', '∝', '≈', '≠', '≤', '≥', '∈', '∉', '⊂', '⊃', '∪', '∩',
        '±', '×', '÷', '⋅', '°', '′', '″', '∠', '⊥', '∥', '∼', '≡', '≜', '≝', '≐',
        
        # 箭头符号
        '←', '→', '↑', '↓', '↔', '↕', '⇐', '⇒', '⇑', '⇓', '⇔', '⇕', '↵', '↩', '↪', '↻', '↺',
        
        # 编辑和文档符号
        '§', '¶', '✎', '✏', '✐', '✁', '✂', '✃', '✄', '✍', '␣', '⌫', '⌧', '⎗', '⎘', '⎙',
        
        # 音乐和多媒体符号
        '♩', '♪', '♫', '♬', '♭', '♮', '♯', '⏮', '⏭', '⏸', '⏹', '⏺', '⏏',
        
        # 中文特殊符号
        '〈', '〉', '《', '》', '「', '」', '『', '』', '【', '】', '〔', '〕', '〖', '〗', '〘', '〙',
        
        # 常用货币符号
        '¢', '€', '£', '¥', '₹', '₽', '₩', '₺', '₽', '₿',
        
        # 商业和法律符号
        '©', '®', '™', '℠', '℗', '℡', '℻'
    ]
    
    # 对特殊符号进行处理（保留原样）- 修复多字节Unicode字符的问题
    symbol_placeholders = {}  # 使用字典存储符号和对应的占位符
    
    for i, symbol in enumerate(special_symbols):
        if symbol in text:
            # 使用索引而不是ord()创建占位符
            placeholder = f"__SYMBOL_{i}_PLACEHOLDER__"
            symbol_placeholders[symbol] = placeholder
            text = text.replace(symbol, placeholder)
    
    # 处理标题
    level = get_heading_level(para)
    if level > 0:
        # 恢复特殊符号
        for symbol, placeholder in symbol_placeholders.items():
            text = text.replace(placeholder, symbol)
        return f"{'#' * level} {text.strip()}"
    
    # 处理列表项
    if is_list_item(para):
        # 恢复特殊符号
        for symbol, placeholder in symbol_placeholders.items():
            text = text.replace(placeholder, symbol)
        return format_list_item(para)
    
    # 处理代码块
    if is_code_block(para):
        # 代码块内不需要转义处理，直接恢复特殊符号
        for symbol, placeholder in symbol_placeholders.items():
            text = text.replace(placeholder, symbol)
        
        # 尝试检测代码语言
        code_language = ""
        code_text = text.strip()
        
        # 根据代码特征判断语言
        if re.search(r'\b(def|class|import|from|print)\b', code_text):
            code_language = "python"
        elif re.search(r'\b(function|var|let|const|require|console\.log)\b', code_text):
            code_language = "javascript"
        elif re.search(r'\b(SELECT|INSERT|UPDATE|DELETE|CREATE|ALTER)\b', code_text, re.IGNORECASE):
            code_language = "sql"
        elif re.search(r'\b(docker|apt|yum|rpm|git|cd|mkdir|tar|curl)\b', code_text):
            code_language = "bash"
        elif re.search(r'<\w+>.*?</w+>|<\w+.*?/>', code_text):
            code_language = "html"
        elif re.search(r'^\s*#include|int\s+main\s*\(', code_text):
            code_language = "cpp"
        # 新增语言检测
        elif re.search(r'\b(package|import java|public class|public static void main)\b', code_text):
            code_language = "java"
        elif re.search(r'\b(using System|namespace|public class|private void)\b', code_text):
            code_language = "csharp"
        elif re.search(r'\b(func|package main|import \(|fmt\.)\b', code_text):
            code_language = "go"
        elif re.search(r'\b(<?php|echo|namespace|use [\w\\]+;)\b', code_text):
            code_language = "php"
        elif re.search(r'\b(fn|let mut|impl|struct|enum|match)\b', code_text):
            code_language = "rust"
        elif re.search(r'\b(library|tidyverse|dplyr|ggplot2|data\.frame)\b', code_text):
            code_language = "r"
        elif re.search(r'(^|\n)[ \t]*(#|//|;)[ \t]*\[?[A-Za-z0-9-_]+\]?[ \t]*:', code_text):
            code_language = "yaml"
        elif re.search(r'\{\s*"[^"]+"\s*:\s*[^{}]+\}', code_text):
            code_language = "json"
        elif re.search(r'\b(module|export|component|ngOnInit|@Input|@Output)\b', code_text):
            code_language = "typescript"
        elif re.search(r'\$\(.*\)|\$\.\w+\(', code_text):
            code_language = "jquery"
        elif re.search(r'<template>|export default {|methods:|computed:', code_text):
            code_language = "vue"
        
        return f"```{code_language}\n{code_text}\n```"
    
    # 处理文本格式（粗体、斜体、链接等）
    formatted_runs = []
    
    for run in para.runs:
        if not run.text:
            continue
        
        run_text = run.text
        
        # 恢复特殊符号
        for symbol, placeholder in symbol_placeholders.items():
            run_text = run_text.replace(placeholder, symbol)
        
        # 处理超链接
        if hasattr(run, 'hyperlink') and run.hyperlink:
            url = run.hyperlink.address
            if url:
                run_text = f"[{run_text}]({url})"
        
        # 处理格式（粗体、斜体）
        is_bold = hasattr(run, 'bold') and run.bold
        is_italic = hasattr(run, 'italic') and run.italic
        
        if is_bold and is_italic:
            run_text = f"***{run_text}***"
        elif is_bold:
            run_text = f"**{run_text}**"
        elif is_italic:
            run_text = f"*{run_text}*"
        
        formatted_runs.append(run_text)
    
    # 将处理后的文本重新组合
    if formatted_runs:
        # 使用处理后的格式化文本块
        return "".join(formatted_runs).strip()
    else:
        # 如果没有runs或处理失败，使用原文本
        # 恢复特殊符号
        for symbol, placeholder in symbol_placeholders.items():
            text = text.replace(placeholder, symbol)
        return text.strip()

def format_table(table):
    """格式化表格为Markdown"""
    md_rows = []
    
    # 确保表格至少有一行
    if len(table.rows) == 0:
        return ""
    
    # 处理表头（第一行）
    header_cells = []
    for cell in table.rows[0].cells:
        header_cells.append(cell.text.strip() or " ")
    md_rows.append("| " + " | ".join(header_cells) + " |")
    
    # 添加分隔行
    md_rows.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
    
    # 处理数据行
    for row in table.rows[1:]:
        row_cells = []
        for cell in row.cells:
            # 确保单元格不为空，避免表格渲染问题
            cell_text = cell.text.strip() or " "
            # 替换表格中的管道符，防止破坏表格结构
            cell_text = cell_text.replace("|", "\\|")
            row_cells.append(cell_text)
        md_rows.append("| " + " | ".join(row_cells) + " |")
    
    return "\n".join(md_rows)

def normalize_heading_levels(md_blocks):
    """标准化标题层级，避免不合理的层级跳跃"""
    normalized_blocks = []
    current_level = 0
    heading_pattern = re.compile(r'^(#+)\s+(.*)$')
    
    for block in md_blocks:
        match = heading_pattern.match(block) if isinstance(block, str) else None
        
        if match:
            hashes, heading_text = match.groups()
            level = len(hashes)
            
            # 第一个标题可以是任何级别
            if current_level == 0:
                current_level = level
            # 不允许直接从低级标题跳到高级标题（如从h1直接跳到h4）
            elif level > current_level + 1:
                # 将层级限制在当前层级+1
                new_level = current_level + 1
                block = f"{'#' * new_level} {heading_text}"
            
            # 更新当前标题级别
            current_level = level if match else current_level
        
        normalized_blocks.append(block)
    
    return normalized_blocks

def insert_images_inline(md_blocks, image_paths, image_dir):
    """根据图片在原文档中的上下文位置精确插入图片引用"""
    # 如果没有图片，直接返回原内容
    if not image_paths:
        return md_blocks
    
    # 创建图片引用列表，包括原始引用ID、文件名和上下文
    image_refs = []
    for i, (ref, path, img_num, context_info) in enumerate(image_paths):
        filename = os.path.basename(path)
        image_refs.append((ref, filename, img_num, context_info))
    
    # 结果列表
    md_with_images = []
    used_images = set()  # 跟踪已插入的图片
    
    # 创建上下文匹配映射
    # 为每个图片计算与每个Markdown块的匹配分数
    matches = []
    
    # 第一步：识别明确的图片引用
    # 查找段落中有明确图片编号的情况，如"图1"、"Figure 2"等
    for i, block in enumerate(md_blocks):
        if not isinstance(block, str):
            continue
            
        # 检查是否有明确的图片编号引用
        fig_match = re.search(r'(图|figure|fig\.)\s*(\d+)', block.lower())
        if fig_match:
            fig_num = int(fig_match.group(2))
            # 找对应编号的图片
            for ref, filename, img_num, _ in image_refs:
                if img_num == fig_num and ref not in used_images:
                    matches.append((i, ref, 150))  # 最高优先级
    
    # 第二步：使用扩展的上下文信息匹配
    for i, block in enumerate(md_blocks):
        if not isinstance(block, str):
            continue
            
        # 对每个图片计算与当前块的匹配分数
        for ref, filename, img_num, context_info in image_refs:
            if not context_info:  # 没有上下文信息则跳过
                continue
                
            # 计算匹配分数
            score = 0
            
            # 检查当前段落文本和图片所在段落的匹配度
            current_para = context_info.get('current_paragraph', '')
            
            if current_para and block:
                # 如果当前块包含图片所在段落的完整文本，给最高分
                if current_para.strip() == block.strip():
                    score += 100
                # 如果当前块包含图片所在段落的文本，给高分
                elif current_para.strip() in block:
                    score += 80
                # 如果有部分文本匹配，根据匹配程度给分
                elif len(current_para) >= 10:
                    # 计算最长公共子串
                    common_text = longest_common_substring(current_para, block)
                    if len(common_text) >= 10:
                        match_ratio = len(common_text) / len(current_para)
                        score += int(60 * match_ratio)
            
            # 字符级匹配 - 检查图片前后的文本
            text_before_image = context_info.get('text_before_image', '')
            text_after_image = context_info.get('text_after_image', '')
            
            # 如果当前块包含了图片前的文本，这是一个很好的插入位置
            if text_before_image and text_before_image.strip() and text_before_image.strip() in block:
                score += 40
                
                # 如果图片前的文本在块的末尾，这是一个完美的插入位置
                if block.strip().endswith(text_before_image.strip()):
                    score += 30
            
            # 检查上下文段落
            context_before = context_info.get('context_paragraphs_before', [])
            context_after = context_info.get('context_paragraphs_after', [])
            
            # 向前查找匹配的上下文段落
            for j, prev_context in enumerate(reversed(context_before)):
                if not prev_context.strip():
                    continue
                    
                # 查找前面的段落是否与当前位置前的块匹配
                for k in range(1, min(6, i+1)):
                    prev_block_index = i - k
                    if prev_block_index < 0:
                        break
                        
                    prev_block = md_blocks[prev_block_index] if isinstance(md_blocks[prev_block_index], str) else ""
                    
                    if prev_context.strip() == prev_block.strip():
                        score += 25 - j*5  # 离当前段落越近，分数越高
                        break
                    elif prev_context.strip() in prev_block:
                        score += 15 - j*3
                        break
                    elif len(prev_context) >= 10:
                        common_text = longest_common_substring(prev_context, prev_block)
                        if len(common_text) >= 10:
                            match_ratio = len(common_text) / len(prev_context)
                            score += int((10 - j*2) * match_ratio)
                            break
            
            # 向后查找匹配的上下文段落
            for j, next_context in enumerate(context_after):
                if not next_context.strip():
                    continue
                    
                # 查找后面的段落是否与当前位置后的块匹配
                for k in range(1, min(6, len(md_blocks)-i)):
                    next_block_index = i + k
                    if next_block_index >= len(md_blocks):
                        break
                        
                    next_block = md_blocks[next_block_index] if isinstance(md_blocks[next_block_index], str) else ""
                    
                    if next_context.strip() == next_block.strip():
                        score += 25 - j*5  # 离当前段落越近，分数越高
                        break
                    elif next_context.strip() in next_block:
                        score += 15 - j*3
                        break
                    elif len(next_context) >= 10:
                        common_text = longest_common_substring(next_context, next_block)
                        if len(common_text) >= 10:
                            match_ratio = len(common_text) / len(next_context)
                            score += int((10 - j*2) * match_ratio)
                            break
            
            # 检查块中是否包含图片指示词
            if re.search(r'(图|figure|image|如图|图片|见图|如下图|示意图|截图|图表|示例|下图|界面|流程图|架构图|结构图)', block.lower()):
                score += 20
            
            # 如果段落以冒号结尾，可能后面跟着图片
            if block.strip().endswith((':', '：')):
                score += 15
            
            # 段落结尾是句号但没有结束词，可能跟着图片
            if block.strip().endswith(('.', '。', '!', '！', '?', '？')):
                score += 5
            
            # 添加到匹配列表，如果分数足够高
            if score >= 20:  # 只保留高于一定分数的匹配
                matches.append((i, ref, score))
    
    # 对匹配按位置和分数排序
    matches.sort(key=lambda x: (x[0], -x[2]))  # 按段落索引升序，分数降序
    
    # 第三步：确保所有图片都有位置
    # 如果有图片没有找到合适的位置，添加合理的位置
    found_refs = set(match[1] for match in matches)
    for ref, filename, img_num, context_info in image_refs:
        if ref not in found_refs:
            # 对于没有找到位置的图片，使用启发式方法找个合适位置
            best_pos = find_best_position_for_image(ref, img_num, md_blocks, image_refs)
            if best_pos is not None:
                matches.append((best_pos, ref, 10))  # 使用较低的分数
    
    # 重新排序匹配
    matches.sort(key=lambda x: (x[0], -x[2]))
    
    # 处理可能的重复：同一位置只保留分数最高的图片
    filtered_matches = []
    pos_map = {}
    
    for pos, ref, score in matches:
        if pos not in pos_map or score > pos_map[pos][1]:
            pos_map[pos] = (ref, score)
    
    for pos, (ref, score) in sorted(pos_map.items()):
        filtered_matches.append((pos, ref, score))
    
    # 合并图片和内容块
    current_match_index = 0
    
    # 处理每个内容块
    for i, block in enumerate(md_blocks):
        md_with_images.append(block)
        
        # 检查当前位置是否需要插入图片
        while current_match_index < len(filtered_matches) and filtered_matches[current_match_index][0] == i:
            pos, ref, score = filtered_matches[current_match_index]
            
            # 获取图片信息
            for r, filename, img_num, _ in image_refs:
                if r == ref and ref not in used_images:
                    md_with_images.append(f"\n![图片{img_num}]({filename})\n")
                    used_images.add(ref)
                    break
                    
            current_match_index += 1
    
    # 确保所有图片都被插入
    # 检查是否有未插入的图片，将它们添加到文档末尾
    remaining_images = [(r, fn, num) for r, fn, num, _ in image_refs if r not in used_images]
    if remaining_images:
        md_with_images.append("\n## 附录：其他图片\n")
        for ref, filename, img_num in remaining_images:
            md_with_images.append(f"\n![图片{img_num}]({filename})\n")
    
    return md_with_images

def find_best_position_for_image(ref, img_num, md_blocks, image_refs):
    """为没有找到匹配位置的图片找一个最佳位置"""
    # 尝试找一个有图片描述词的段落
    for i, block in enumerate(md_blocks):
        if not isinstance(block, str):
            continue
            
        # 检查是否包含图片描述词但没有具体编号
        if re.search(r'(图|figure|image|如图|图片|见图|如下图|示意图|截图|图表|示例|下图)', block.lower()) and not re.search(r'图\s*\d+', block.lower()):
            return i
    
    # 如果没有找到，查找段落结束有冒号的位置
    for i, block in enumerate(md_blocks):
        if not isinstance(block, str):
            continue
            
        if block.strip().endswith((':', '：')):
            return i
    
    # 如果没有找到合适位置，尝试找到标题后的位置
    heading_positions = []
    for i, block in enumerate(md_blocks):
        if not isinstance(block, str):
            continue
            
        if block.startswith('#'):
            heading_positions.append(i)
    
    # 如果找到了标题，将图片放在标题后
    if heading_positions:
        # 将图片放在第一个标题后的位置
        for i in heading_positions:
            if i < len(md_blocks) - 1:
                return i  # 返回标题的位置，图片会被插入到标题后面
    
    # 找不到合适位置，返回文档靠前的位置（不放在最后，避免所有图片都堆积到附录）
    return min(5, len(md_blocks) - 1) if len(md_blocks) > 5 else 0

def longest_common_substring(s1, s2):
    """查找两个字符串的最长公共子串"""
    # 简化版本，优化性能
    if not s1 or not s2:
        return ""
    
    # 对于长字符串，只检查前200个字符
    s1 = s1[:200] if len(s1) > 200 else s1
    s2 = s2[:200] if len(s2) > 200 else s2
    
    # 初始化DP表
    m, n = len(s1), len(s2)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    
    # 填充DP表
    max_length = 0
    end_pos = 0
    
    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if s1[i-1] == s2[j-1]:
                dp[i][j] = dp[i-1][j-1] + 1
                if dp[i][j] > max_length:
                    max_length = dp[i][j]
                    end_pos = i
    
    # 提取最长公共子串
    if max_length == 0:
        return ""
    return s1[end_pos - max_length:end_pos]

def convert_docx_to_md(docx_path, output_path, image_dir=None):
    """将docx文件转换为markdown格式"""
    try:
        logger.info(f"开始转换DOCX文件: {docx_path}")
        
        if image_dir is None:
            # 默认将图片放在与输出文件相同的目录中
            image_dir = os.path.dirname(output_path)
        
        # 确保输出目录存在
        output_dirname = os.path.dirname(output_path)
        if output_dirname:
            os.makedirs(output_dirname, exist_ok=True)
        
        # 确保图片目录存在
        os.makedirs(image_dir, exist_ok=True)
        
        # 加载文档
        doc = docx.Document(docx_path)
        md_blocks = []
        
        # 提取图片
        logger.info("提取文档中的图片和位置信息")
        image_paths = extract_images_from_docx(docx_path, image_dir)
        logger.info(f"共提取了 {len(image_paths)} 张图片")
        
        # 提取文档标题，优先使用Title样式
        title = None
        for para in doc.paragraphs:
            if para.style and hasattr(para.style, 'name'):
                style_name = para.style.name
                if 'Title' in style_name or '标题' in style_name or style_name == 'Title':
                    title = para.text.strip()
                    break
        
        # 如果没有找到标题样式，尝试查找文档第一个段落是否是标题
        if not title and doc.paragraphs:
            first_para = doc.paragraphs[0]
            if get_heading_level(first_para) == 1:
                title = first_para.text.strip()
        
        # 提取目录结构
        logger.info("提取文档结构")
        toc = extract_toc(doc)
        
        # 如果有标题，添加到开始
        if title:
            md_blocks.append(f"# {title}\n")
        
        # 添加目录（仅当文档较长且有多个标题时）
        if toc and len(toc) > 3:
            md_blocks.append(generate_toc_md(toc))
        
        # 使用迭代器处理文档的所有块（段落和表格），保持顺序
        logger.info("处理文档内容")
        for item in iter_block_items(doc):
            if isinstance(item, Paragraph):
                # 处理段落
                md_text = format_paragraph(item)
                if md_text:
                    md_blocks.append(md_text)
            elif isinstance(item, Table):
                # 处理表格
                md_table = format_table(item)
                if md_table:
                    md_blocks.append(md_table)
        
        # 标准化标题层级
        logger.info("标准化标题层级")
        md_blocks = normalize_heading_levels(md_blocks)
        
        # 基于上下文信息精准插入图片
        logger.info("根据上下文精准插入图片引用")
        md_blocks = insert_images_inline(md_blocks, image_paths, image_dir)
        
        # 合并所有块并插入适当的空行
        md_content = []
        prev_block_type = None
        
        for block in md_blocks:
            if not isinstance(block, str):
                continue
                
            block = block.strip()
            current_block_type = None
            
            # 识别块类型以决定间距
            if block.startswith('#'):
                current_block_type = 'heading'
            elif block.startswith('```'):
                current_block_type = 'code'
            elif block.startswith('|') and '|' in block[1:]:
                current_block_type = 'table'
            elif block.startswith('!['):
                current_block_type = 'image'
            elif re.match(r'^\s*[-*+]\s', block):
                current_block_type = 'unordered_list'
            elif re.match(r'^\s*\d+\.\s', block):
                current_block_type = 'ordered_list'
            else:
                current_block_type = 'paragraph'
            
            # 根据前后块类型添加空行
            if prev_block_type and prev_block_type != current_block_type:
                # 标题前要有额外空行
                if current_block_type == 'heading':
                    md_content.append('')
                # 代码块前后要有空行
                elif prev_block_type == 'code' or current_block_type == 'code':
                    md_content.append('')
                # 表格前后要有空行
                elif prev_block_type == 'table' or current_block_type == 'table':
                    md_content.append('')
                # 图片前后要有空行
                elif prev_block_type == 'image' or current_block_type == 'image':
                    md_content.append('')
                # 列表项和段落间要有空行
                elif (prev_block_type in ['paragraph', 'unordered_list', 'ordered_list'] and 
                      current_block_type in ['paragraph', 'unordered_list', 'ordered_list'] and
                      prev_block_type != current_block_type):
                    md_content.append('')
            
            md_content.append(block)
            prev_block_type = current_block_type
        
        # 最终内容
        final_md_content = '\n'.join(md_content)
        
        # 写入markdown文件
        logger.info(f"写入Markdown文件: {output_path}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_md_content)
        
        # 返回图片路径列表，以便后续处理
        logger.info("DOCX转换完成")
        return [os.path.join(image_dir, os.path.basename(path)) for _, path, _, _ in image_paths]
    
    except Exception as e:
        logger.error(f"转换DOCX文件时出错: {str(e)}", exc_info=True)
        raise 