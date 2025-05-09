"""
本地CYRUS转换方法实现
在无法从GitHub安装官方库时使用

此实现提供了基本的文档转换功能，模拟CYRUS-STUDIO/docx2markdown的行为
"""

import os
import logging
import re
import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
from io import BytesIO

logger = logging.getLogger(__name__)

def iter_block_items(parent):
    """按顺序迭代文档中的所有段落和表格"""
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    else:
        raise ValueError("不支持的父元素类型")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_images_from_docx(doc_path, image_dir):
    """从docx文件中提取图片"""
    doc = docx.Document(doc_path)
    image_paths = []
    
    # 确保图片目录存在
    os.makedirs(image_dir, exist_ok=True)
    
    # 提取图片
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            image_data = rel.target_part.blob
            
            # 尝试确定图片格式
            try:
                img = Image.open(BytesIO(image_data))
                ext = img.format.lower()
            except:
                ext = "png"  # 默认使用png
            
            image_filename = f"image_{image_count}.{ext}"
            image_path = os.path.join(image_dir, image_filename)
            
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            image_paths.append(image_path)
    
    return image_paths

def process_paragraph(paragraph):
    """处理文档段落"""
    if not paragraph.text.strip():
        return ""

    # 检查是否为标题
    if paragraph.style and "heading" in paragraph.style.name.lower():
        level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
        return f"{'#' * level} {paragraph.text.strip()}\n"
        
    text = paragraph.text.strip()
    
    # 处理基本格式
    result = []
    for run in paragraph.runs:
        run_text = run.text
        
        # 加粗
        if run.bold:
            run_text = f"**{run_text}**"
            
        # 斜体
        if run.italic:
            run_text = f"*{run_text}*"
            
        result.append(run_text)
    
    formatted_text = "".join(result)
    
    # 处理列表
    if re.match(r'^\d+[\.\)]\s', text):
        return f"{text}\n"
    elif re.match(r'^[\*\-\•]\s', text):
        return f"{text}\n"
    
    return f"{formatted_text}\n"

def process_table(table):
    """处理文档表格"""
    result = []
    
    # 表头
    header = []
    for cell in table.rows[0].cells:
        header.append(cell.text.strip())
    
    result.append("| " + " | ".join(header) + " |")
    result.append("| " + " | ".join(["---"] * len(header)) + " |")
    
    # 表格内容
    for row in table.rows[1:]:
        cells = []
        for cell in row.cells:
            cells.append(cell.text.strip())
        result.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(result) + "\n"

def convert_docx_to_md_local_cyrus(docx_path, output_path, image_dir=None):
    """
    本地实现的CYRUS风格文档转换
    
    Args:
        docx_path (str): DOCX文件路径
        output_path (str): 输出的Markdown文件路径
        image_dir (str, optional): 图片保存目录
    
    Returns:
        list: 包含图片路径的列表
    """
    try:
        logger.info(f"使用本地CYRUS实现转换DOCX文件: {docx_path}")
        
        if image_dir is None:
            # 默认将图片放在输出文件所在目录下的images文件夹中
            image_dir = os.path.join(os.path.dirname(output_path), 'images')
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        os.makedirs(image_dir, exist_ok=True)
        
        # 提取图片
        image_paths = extract_images_from_docx(docx_path, image_dir)
        logger.info(f"提取了 {len(image_paths)} 张图片")
        
        # 处理文档内容
        doc = docx.Document(docx_path)
        md_content = []
        
        # 处理文档标题
        if doc.paragraphs and doc.paragraphs[0].style and "title" in doc.paragraphs[0].style.name.lower():
            md_content.append(f"# {doc.paragraphs[0].text.strip()}\n")
        
        # 处理文档内容
        for item in iter_block_items(doc):
            if isinstance(item, Paragraph):
                paragraph_md = process_paragraph(item)
                if paragraph_md:
                    md_content.append(paragraph_md)
            elif isinstance(item, Table):
                table_md = process_table(item)
                if table_md:
                    md_content.append(table_md)
        
        # 处理图片引用
        for i, img_path in enumerate(image_paths):
            img_filename = os.path.basename(img_path)
            img_rel_path = os.path.join('images', img_filename)
            md_content.append(f"![图片 {i+1}]({img_rel_path})\n")
        
        # 写入输出文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(md_content))
        
        logger.info(f"本地CYRUS实现转换完成: {output_path}")
        return image_paths
    
    except Exception as e:
        logger.error(f"本地CYRUS实现转换出错: {str(e)}", exc_info=True)
        return [] 