import os
import io
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_to_md import convert_docx_to_md

def create_test_docx_with_lists(file_path):
    """创建一个包含各种列表类型的测试DOCX文件"""
    print(f"创建列表测试DOCX文件: {file_path}")
    
    doc = Document()
    
    # 添加文档标题
    title = doc.add_paragraph("列表格式测试文档")
    title.style = doc.styles['Title']
    
    # 添加简介
    doc.add_paragraph("本文档用于测试各种列表格式的转换效果，包括有序列表、无序列表和嵌套列表。")
    
    # 1. 无序列表测试
    h1 = doc.add_heading('1. 无序列表测试', level=1)
    p = doc.add_paragraph("以下是不同类型的无序列表:")
    
    # 常规无序列表
    doc.add_paragraph("常规无序列表项1", style='List Bullet')
    doc.add_paragraph("常规无序列表项2", style='List Bullet')
    doc.add_paragraph("常规无序列表项3", style='List Bullet')
    
    # 使用符号的无序列表
    p = doc.add_paragraph("• 使用圆点符号的列表项1")
    p = doc.add_paragraph("• 使用圆点符号的列表项2")
    p = doc.add_paragraph("• 使用圆点符号的列表项3")
    
    # 使用破折号的无序列表
    p = doc.add_paragraph("- 使用破折号的列表项1")
    p = doc.add_paragraph("- 使用破折号的列表项2")
    p = doc.add_paragraph("- 使用破折号的列表项3")
    
    # 2. 有序列表测试
    h1 = doc.add_heading('2. 有序列表测试', level=1)
    p = doc.add_paragraph("以下是不同类型的有序列表:")
    
    # 数字加点号的有序列表
    doc.add_paragraph("1. 第一条", style='List Number')
    doc.add_paragraph("2. 第二条", style='List Number')
    doc.add_paragraph("3. 第三条", style='List Number')
    
    # 数字加括号的有序列表
    p = doc.add_paragraph("1) 第一条带括号")
    p = doc.add_paragraph("2) 第二条带括号")
    p = doc.add_paragraph("3) 第三条带括号")
    
    # 数字加中文顿号的有序列表
    p = doc.add_paragraph("1、第一条中文格式")
    p = doc.add_paragraph("2、第二条中文格式")
    p = doc.add_paragraph("3、第三条中文格式")
    
    # 3. 嵌套列表测试
    h1 = doc.add_heading('3. 嵌套列表测试', level=1)
    p = doc.add_paragraph("以下是嵌套列表样例:")
    
    # 嵌套无序列表
    doc.add_paragraph("无序一级项1", style='List Bullet')
    doc.add_paragraph("无序二级项1.1", style='List Bullet 2')
    doc.add_paragraph("无序二级项1.2", style='List Bullet 2')
    doc.add_paragraph("无序一级项2", style='List Bullet')
    doc.add_paragraph("无序二级项2.1", style='List Bullet 2')
    doc.add_paragraph("无序三级项2.1.1", style='List Bullet 3')
    
    # 嵌套有序列表
    doc.add_paragraph("1. 有序一级项1", style='List Number')
    doc.add_paragraph("a. 有序二级项1.1", style='List Number 2')
    doc.add_paragraph("b. 有序二级项1.2", style='List Number 2')
    doc.add_paragraph("2. 有序一级项2", style='List Number')
    doc.add_paragraph("a. 有序二级项2.1", style='List Number 2')
    doc.add_paragraph("i. 有序三级项2.1.1", style='List Number 3')
    
    # 混合嵌套列表
    doc.add_paragraph("• 混合列表一级项1")
    p = doc.add_paragraph("  1. 混合列表二级项1.1")
    p.paragraph_format.left_indent = Pt(24)
    p = doc.add_paragraph("  2. 混合列表二级项1.2")
    p.paragraph_format.left_indent = Pt(24)
    doc.add_paragraph("• 混合列表一级项2")
    p = doc.add_paragraph("  1. 混合列表二级项2.1")
    p.paragraph_format.left_indent = Pt(24)
    p = doc.add_paragraph("    • 混合列表三级项2.1.1")
    p.paragraph_format.left_indent = Pt(48)
    
    # 保存文档
    doc.save(file_path)
    print(f"列表测试DOCX文件已保存: {file_path}")

def test_lists_conversion():
    """测试列表格式的DOCX到Markdown的转换功能"""
    test_docx = 'test_lists_document.docx'
    output_dir = 'outputs'
    output_md = os.path.join(output_dir, 'test_lists_output.md')
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 清理旧的转换结果
    if os.path.exists(output_md):
        os.remove(output_md)
    
    # 清理旧的图片目录
    image_dir = os.path.join(output_dir, 'test_lists_output_images')
    if os.path.exists(image_dir):
        import shutil
        shutil.rmtree(image_dir)
    
    # 创建测试文档
    create_test_docx_with_lists(test_docx)
    
    # 执行转换
    print(f"\n开始转换列表测试DOCX到Markdown...")
    convert_docx_to_md(test_docx, output_md)
    
    # 检查结果
    if os.path.exists(output_md):
        print(f"转换成功！已生成Markdown文件: {output_md}")
        
        # 显示Markdown内容
        with open(output_md, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        print("\nMarkdown内容预览:")
        print("-" * 40)
        print(md_content)
        print("-" * 40)
        
    else:
        print(f"转换失败！未生成Markdown文件。")
    
    # 清理测试文件
    try:
        if os.path.exists(test_docx):
            os.remove(test_docx)
    except:
        pass

if __name__ == "__main__":
    try:
        test_lists_conversion()
        print("\n测试完成。")
    except Exception as e:
        print(f"\n测试过程中出错: {e}")
        import traceback
        traceback.print_exc() 