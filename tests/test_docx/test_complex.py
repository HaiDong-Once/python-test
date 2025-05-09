import os
import io
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_to_md import convert_docx_to_md

def create_test_image(filename, width=400, height=300, color=(0, 120, 212)):
    """创建一个测试图片"""
    img = Image.new('RGB', (width, height), color=color)
    img.save(filename)
    return filename

def create_complex_test_docx(file_path):
    """创建一个复杂的测试DOCX文件，包含标题、格式化文本、表格和图片"""
    print(f"创建复杂测试DOCX文件: {file_path}")
    
    doc = Document()
    
    # 添加文档标题
    title = doc.add_paragraph("有道云笔记转Markdown示例文档")
    title.style = doc.styles['Title']
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加多级标题结构（目录）
    h1 = doc.add_heading('1. 文档简介', level=1)
    p = doc.add_paragraph('这是一个用于测试DOCX到Markdown转换功能的示例文档。本文档包含了不同级别的标题、格式化文本、表格、代码块和图片。')
    
    h1 = doc.add_heading('2. 文本格式化', level=1)
    p = doc.add_paragraph('Markdown支持多种文本格式化，包括 ')
    p.add_run('粗体').bold = True
    p.add_run('、')
    p.add_run('斜体').italic = True
    p.add_run(' 和 ')
    run = p.add_run('粗斜体')
    run.bold = True
    run.italic = True
    p.add_run('。')
    
    # 添加二级标题
    h2 = doc.add_heading('2.1 链接格式', level=2)
    p = doc.add_paragraph('Markdown中的链接格式为 ')
    run = p.add_run('链接文本')
    run.hyperlink = doc.part.relate_to('https://example.com', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    p.add_run('。')
    
    # 添加三级标题
    h3 = doc.add_heading('2.1.1 图片格式', level=3)
    p = doc.add_paragraph('Markdown中的图片格式类似于链接，但前面有一个感叹号。')
    
    # 创建并添加测试图片
    img_path = create_test_image('test_image.png')
    doc.add_picture(img_path, width=Inches(4))
    
    h1 = doc.add_heading('3. 表格示例', level=1)
    p = doc.add_paragraph('下面是一个简单的表格示例：')
    
    # 添加表格
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '序号'
    header_cells[1].text = '项目'
    header_cells[2].text = '描述'
    
    # 填充表格数据
    data = [
        (1, '标题', 'Markdown使用#表示标题，#号越多级别越低'),
        (2, '列表', 'Markdown使用-或*表示无序列表，数字加.表示有序列表'),
        (3, '代码块', 'Markdown使用```包裹代码块')
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i+1].cells
        row[0].text = str(row_data[0])
        row[1].text = row_data[1]
        row[2].text = row_data[2]
    
    h1 = doc.add_heading('4. 代码块示例', level=1)
    p = doc.add_paragraph('下面是一个Python代码块示例：')
    
    # 添加代码块段落
    code = doc.add_paragraph('''def hello_world():
    """打印Hello World"""
    print("Hello, World!")
    
    if __name__ == "__main__":
        hello_world()''')
    code.style = doc.styles['No Spacing']
    
    # 尽量设置代码块样式
    for run in code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    
    h1 = doc.add_heading('5. 结论', level=1)
    p = doc.add_paragraph('通过本文档的示例，我们展示了Markdown格式的主要特性，以及如何从DOCX文档正确转换为Markdown格式。')
    
    # 保存文档
    doc.save(file_path)
    print(f"复杂测试DOCX文件已保存: {file_path}")
    return img_path

def test_complex_conversion():
    """测试复杂DOCX到Markdown的转换功能"""
    test_docx = 'complex_test_document.docx'
    output_dir = 'outputs'
    output_md = os.path.join(output_dir, 'complex_test_output.md')
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 清理旧的转换结果
    if os.path.exists(output_md):
        os.remove(output_md)
    
    # 清理旧的图片目录
    image_dir = os.path.join(output_dir, 'complex_test_output_images')
    if os.path.exists(image_dir):
        import shutil
        shutil.rmtree(image_dir)
    
    # 创建测试文档
    img_path = create_complex_test_docx(test_docx)
    
    # 执行转换
    print(f"\n开始转换复杂DOCX到Markdown...")
    convert_docx_to_md(test_docx, output_md)
    
    # 检查结果
    if os.path.exists(output_md):
        print(f"转换成功！已生成Markdown文件: {output_md}")
        
        # 显示Markdown内容
        with open(output_md, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        print("\nMarkdown内容预览:")
        print("-" * 40)
        print(md_content[:1000] + ("..." if len(md_content) > 1000 else ""))
        print("-" * 40)
        
        # 检查图片目录
        if os.path.exists(image_dir):
            image_files = os.listdir(image_dir)
            print(f"\n图片目录: {image_dir}")
            print(f"图片文件数量: {len(image_files)}")
            if image_files:
                print("图片文件列表:")
                for img in image_files:
                    print(f"  - {img}")
    else:
        print(f"转换失败！未生成Markdown文件。")
    
    # 清理临时文件
    try:
        if os.path.exists(img_path):
            os.remove(img_path)
    except:
        pass

if __name__ == "__main__":
    try:
        test_complex_conversion()
        print("\n测试完成。")
    except Exception as e:
        print(f"\n测试过程中出错: {e}")
        import traceback
        traceback.print_exc() 