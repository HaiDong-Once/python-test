import os
import docx
from utils.docx_to_md import convert_docx_to_md

def create_simple_test_docx():
    """创建一个简单的测试文档，包含列表和代码"""
    doc = docx.Document()
    
    # 添加标题
    doc.add_heading('列表和代码测试', 0)
    
    # 无序列表
    doc.add_paragraph('无序列表:', style='Normal')
    doc.add_paragraph('第一项', style='List Bullet')
    doc.add_paragraph('第二项', style='List Bullet')
    doc.add_paragraph('第三项', style='List Bullet')
    
    # 有序列表
    doc.add_paragraph('有序列表:', style='Normal')
    doc.add_paragraph('第一条', style='List Number')
    doc.add_paragraph('第二条', style='List Number')
    doc.add_paragraph('第三条', style='List Number')
    
    # 嵌套列表
    doc.add_paragraph('嵌套列表:', style='Normal')
    doc.add_paragraph('主列表项1', style='List Bullet')
    doc.add_paragraph('子列表项1.1', style='List Bullet 2')
    doc.add_paragraph('子列表项1.2', style='List Bullet 2')
    doc.add_paragraph('主列表项2', style='List Bullet')
    
    # 代码块
    doc.add_paragraph('代码示例:', style='Normal')
    code = doc.add_paragraph('def hello_world():\n    print("Hello, World!")', style='Normal')
    
    # 保存文档
    test_file = 'test_lists_simple.docx'
    doc.save(test_file)
    print(f"创建测试文档: {test_file}")
    return test_file

def test_conversion():
    """测试文档转换"""
    # 创建测试文档
    test_docx = create_simple_test_docx()
    
    # 转换为Markdown
    output_dir = 'outputs'
    os.makedirs(output_dir, exist_ok=True)
    
    output_md = os.path.join(output_dir, 'test_lists_simple.md')
    print(f"开始转换文档: {test_docx} -> {output_md}")
    
    convert_docx_to_md(test_docx, output_md)
    
    # 显示转换结果
    if os.path.exists(output_md):
        print(f"转换成功! 内容:")
        with open(output_md, 'r', encoding='utf-8') as f:
            print(f.read())
    else:
        print("转换失败!")

if __name__ == "__main__":
    test_conversion() 